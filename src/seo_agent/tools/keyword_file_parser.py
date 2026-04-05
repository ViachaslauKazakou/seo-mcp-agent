"""Parse keyword lists from uploaded files (Excel, CSV, text, Word)."""

import csv
import io
import logging
import re
from collections import Counter
from typing import IO

logger = logging.getLogger(__name__)

# Maximum keywords extracted from a single file upload
MAX_KEYWORDS = 5000

# Candidates for single-character delimiters in plain-text files
_CANDIDATE_DELIMITERS = [",", ";", "\t", "|", ":", "  ", " "]


def _clean_keyword(raw: str) -> str:
    """Strip whitespace and invisible characters; return empty string if meaningless."""
    cleaned = raw.strip().strip("\u200b\ufeff\xa0")
    # Drop values that are purely numeric or single-character noise
    if len(cleaned) < 2:
        return ""
    return cleaned


def _detect_delimiter(text: str) -> str:
    """Return the most likely delimiter by counting repeating occurrences per line."""
    lines = [l for l in text.splitlines() if l.strip()][:50]  # sample first 50 lines
    if not lines:
        return "\n"

    scores: dict[str, int] = {}
    for line in lines:
        for delim in _CANDIDATE_DELIMITERS:
            count = line.count(delim)
            if count > 0:
                scores[delim] = scores.get(delim, 0) + count

    if not scores:
        return "\n"  # one keyword per line

    # Prefer the delimiter that appears consistently across the most lines rather
    # than just having the highest raw count (spaces would otherwise always win).
    consistency: dict[str, int] = {}
    for delim in scores:
        lines_with = sum(1 for l in lines if l.count(delim) > 0)
        consistency[delim] = lines_with

    best = max(consistency, key=lambda d: (consistency[d], scores[d]))

    # Reject single-space delimiter when the text looks like multi-word keywords
    # (i.e., average words per line > 2 and the winning delimiter is a space).
    if best == " ":
        avg_words = sum(len(l.split()) for l in lines) / len(lines)
        if avg_words > 2.5:
            # Fall back to newline — each line is a phrase
            return "\n"

    return best


def _split_by_delimiter(text: str, delimiter: str) -> list[str]:
    """Split text by delimiter, supporting multi-char delimiters (e.g. double space)."""
    if delimiter == "\n":
        parts = text.splitlines()
    elif len(delimiter) > 1:
        parts = re.split(re.escape(delimiter), text)
    else:
        parts = text.split(delimiter)
    return parts


def _extract_from_text(text: str) -> list[str]:
    """Auto-detect delimiter and return cleaned keyword list."""
    delimiter = _detect_delimiter(text)
    logger.debug("Detected delimiter: %r", delimiter)
    parts = _split_by_delimiter(text, delimiter)

    # For newline-separated data, each line may itself contain comma/semicolon lists
    if delimiter == "\n":
        expanded: list[str] = []
        for part in parts:
            sub_delim = _detect_delimiter(part) if len(part) > 5 else None
            if sub_delim and sub_delim != "\n" and part.count(sub_delim) > 0:
                expanded.extend(part.split(sub_delim))
            else:
                expanded.append(part)
        parts = expanded

    keywords = []
    for raw in parts:
        kw = _clean_keyword(raw)
        if kw:
            keywords.append(kw)

    return keywords[:MAX_KEYWORDS]


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------


def parse_txt(content: bytes, encoding: str = "utf-8") -> list[str]:
    """Parse plain text file and extract keywords."""
    try:
        text = content.decode(encoding)
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="replace")
    return _extract_from_text(text)


def parse_csv(content: bytes, encoding: str = "utf-8") -> list[str]:
    """Parse CSV file. Takes the first non-empty value from each row."""
    try:
        text = content.decode(encoding)
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="replace")

    # Try csv.Sniffer first; fall back to auto-detect
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        has_header = csv.Sniffer().has_header(sample)
    except csv.Error:
        dialect = csv.excel  # type: ignore[assignment]
        has_header = False

    reader = csv.reader(io.StringIO(text), dialect)
    rows = list(reader)

    if has_header and rows:
        rows = rows[1:]

    keywords: list[str] = []
    for row in rows:
        for cell in row:
            kw = _clean_keyword(cell)
            if kw:
                keywords.append(kw)
                break  # take first meaningful cell per row

    # If all rows had only one column → might be comma-delimited phrases in cells
    if not keywords and rows:
        flat_text = "\n".join(",".join(r) for r in rows)
        keywords = _extract_from_text(flat_text)

    return keywords[:MAX_KEYWORDS]


def parse_xlsx(content: bytes) -> list[str]:
    """Parse Excel xlsx/xlsm file. Collects all non-empty string cells."""
    import openpyxl  # lazy import – not always needed

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    keywords: list[str] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for cell in row:
                if cell is None:
                    continue
                kw = _clean_keyword(str(cell))
                if kw:
                    keywords.append(kw)
        if len(keywords) >= MAX_KEYWORDS:
            break
    wb.close()
    return keywords[:MAX_KEYWORDS]


def parse_docx(content: bytes) -> list[str]:
    """Parse Word docx file. Extracts keywords from paragraphs and table cells."""
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(content))
    texts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(text)

    combined = "\n".join(texts)
    return _extract_from_text(combined)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse_keyword_file(filename: str, content: bytes) -> list[str]:
    """Dispatch to the correct parser based on file extension.

    Returns a deduplicated list of keyword strings (lowercase, stripped).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("xlsx", "xlsm", "xls"):
        if ext == "xls":
            # xlrd handles old xls; fall back to treating as CSV-ish if not available
            try:
                import xlrd  # type: ignore[import]

                wb = xlrd.open_workbook(file_contents=content)
                keywords: list[str] = []
                for sheet in wb.sheets():
                    for row_idx in range(sheet.nrows):
                        for col_idx in range(sheet.ncols):
                            cell = sheet.cell_value(row_idx, col_idx)
                            if cell:
                                kw = _clean_keyword(str(cell))
                                if kw:
                                    keywords.append(kw)
                return _dedup(keywords)[:MAX_KEYWORDS]
            except ImportError:
                logger.warning("xlrd not installed; attempting openpyxl for .xls file")
                return _dedup(parse_xlsx(content))
        return _dedup(parse_xlsx(content))

    elif ext == "csv":
        return _dedup(parse_csv(content))

    elif ext in ("txt", "text", "md"):
        return _dedup(parse_txt(content))

    elif ext == "docx":
        return _dedup(parse_docx(content))

    elif ext == "doc":
        # .doc (legacy Word binary) — try antiword-style approach via python-docx2txt
        try:
            import docx2txt  # type: ignore[import]

            text = docx2txt.process(io.BytesIO(content))
            return _dedup(_extract_from_text(text))
        except ImportError:
            logger.warning("docx2txt not installed; cannot parse .doc format")
            return []

    else:
        # Unknown extension — try plain-text heuristic
        logger.warning("Unknown file extension %r, attempting plain-text parse", ext)
        return _dedup(parse_txt(content))


def _dedup(keywords: list[str]) -> list[str]:
    """Deduplicate preserving order (case-insensitive)."""
    seen: set[str] = set()
    result: list[str] = []
    for kw in keywords:
        key = kw.lower()
        if key not in seen:
            seen.add(key)
            result.append(kw)
    return result
